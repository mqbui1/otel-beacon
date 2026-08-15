"""
Generate otel-beacon Anomaly Detection Engine — Engineering Reference (DOCX).
Run: python3 docs/gen_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Style helpers ────────────────────────────────────────────────────────────
DARK   = RGBColor(0x1a, 0x1a, 0x2e)
TEAL   = RGBColor(0x0a, 0x84, 0x8a)
ORANGE = RGBColor(0xd9, 0x73, 0x06)
GRAY   = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0xc0, 0x20, 0x20)

def set_heading(para, text, level=1):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(20)
        run.font.color.rgb = TEAL
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = DARK
    elif level == 3:
        run.font.size  = Pt(12)
        run.font.color.rgb = TEAL
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = ORANGE

def h(level, text):
    p = doc.add_heading('', level=min(level, 9))
    set_heading(p, text, level)
    return p

def p(text='', bold_prefix=None, mono=False):
    para = doc.add_paragraph()
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        para.add_run(' ')
    if text:
        r = para.add_run(text)
        if mono:
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
    para.paragraph_format.space_after = Pt(4)
    return para

def code(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    r = para.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1e, 0x60, 0x91)
    # light gray shading
    rPr = r._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F4F8')
    rPr.append(shd)
    return para

def bullet(text, bold_prefix=None, indent=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        para.add_run(' ')
    if text:
        para.add_run(text)
    para.paragraph_format.space_after = Pt(2)
    return para

def table_header(tbl, headers):
    row = tbl.rows[0]
    for i, h_text in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '0A848A')
        cell._tc.tcPr.append(shd)

def add_table_row(tbl, values):
    row = tbl.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(v)
        r.font.size = Pt(9)
    return row

def hr():
    p_el = doc.add_paragraph()
    p_el.paragraph_format.space_before = Pt(4)
    p_el.paragraph_format.space_after  = Pt(4)
    pPr = p_el._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('otel-beacon')
title_run.bold = True
title_run.font.size = Pt(32)
title_run.font.color.rgb = TEAL

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_para.add_run('Anomaly Detection Engine — Engineering Reference')
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = GRAY

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run(f'Generated {datetime.date.today().isoformat()}  ·  Internal Engineering Document')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h(1, '1. Architecture Overview')
p('otel-beacon is a self-hosted OpenTelemetry backend written in Go. It receives '
  'OTLP/gRPC (port 4317) and OTLP/HTTP (port 4318) traffic, persists everything to '
  'SQLite (or ClickHouse), and runs a multi-layered anomaly detection pipeline entirely '
  'in-process. The admin + query API is served on port 8080.')

h(2, '1.1  Data Flow')
p('Each telemetry signal type has its own async processing path:')
bullet('Spans → spanCh (buffered channel, cap 10 000) → spanWorker → FlushSpans + UpsertEntities')
bullet('Metrics → metricCh → metricsWorker → anomaly detection inline → FlushMetrics')
bullet('Logs → logCh → logsWorker → FlushLogs + UpsertEntities')
bullet('GenAI spans (gen_ai.*) → genaiCh → genaiWorker → FlushGenAISpans + eval queue')
p('Workers batch items and flush every 500 ms or when a batch of 500 is full, '
  'whichever comes first. Failed flushes are retried with exponential backoff (max 5 attempts, '
  'starting at 100 ms, doubling each attempt).')

h(2, '1.2  Background Workers (goroutines)')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Worker', 'Interval', 'Responsibility'])
rows = [
    ('spanWorker',            '500 ms flush',        'Batch-flush spans; extract + upsert entity records'),
    ('metricsWorker',         '500 ms flush',        'Flush metrics; run inline metric anomaly detector'),
    ('logsWorker',            '500 ms flush',        'Flush logs; extract + upsert entity records'),
    ('genaiWorker',           '500 ms flush',        'Flush GenAI spans; route to eval + session queues'),
    ('topologyWorker',        '2 min',               'RefreshTopology — rebuild service_topology table from spans'),
    ('fingerprintWorker',     '5 min (2 min delay)', 'Compute trace structural fingerprints; detect trace_drift'),
    ('errorSignatureWorker',  '5 min (2 min delay)', 'Build per-service error signatures; detect spikes'),
    ('spanRateWorker',        '5 min + 30 s fast',   'Error rate + P95 latency detection (two parallel detectors)'),
    ('callGraphDriftWorker',  '2 min',               'Detect new topology edges → callgraph_drift anomaly'),
    ('retentionWorker',       '1 hour',              'Delete rows older than RETENTION_DAYS (default 30)'),
    ('StartMissingSvcChecker','10 s',                'Detect services silent for >45 s → missing_service anomaly'),
    ('StartCorrelator',       '60 s',                'Cross-signal correlation → correlated_incident'),
    ('StartEvalWorker',       'drain queue',         'LLM-as-judge eval via AWS Bedrock (or heuristic fallback)'),
    ('StartSessionEvalWorker','drain queue',         'Session-level quality scoring'),
]
for r in rows:
    add_table_row(tbl, r)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ENTITY FRAMEWORK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '2. Entity Framework')
p('Entities are the canonical identifiers that correlate all telemetry signals. '
  'Every span, metric, log, and anomaly is associated with an entity_id resolved at ingest time.')

h(2, '2.1  Entity Resolution (extractEntityID)')
p('The entity ID is resolved from the OTLP resource attributes JSON stored with each telemetry record:')
bullet('Primary key: service.name attribute (preferred for all microservices)')
bullet('Fallback: host.name (used by infrastructure exporters that report node-level metrics/logs without a service name)')
bullet('If neither is present, entity_id is empty and the record is stored without entity linkage')

code('func extractEntityID(resJSON string) string {\n'
     '    // 1. Try service.name\n'
     '    if svc := attrs["service.name"]; svc != "" { return svc }\n'
     '    // 2. Fall back to host.name\n'
     '    if host := attrs["host.name"]; host != "" { return host }\n'
     '    return ""\n'
     '}')

h(2, '2.2  Entity Record Schema (EntityRow)')
p('Entities are upserted into the entities table on every span/metric/log flush:')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Field', 'Type', 'Description'])
for r in [
    ('entity_type',  'TEXT',     '"service" or "host"'),
    ('entity_id',    'TEXT PK',  'Canonical name (service.name or host.name)'),
    ('environment',  'TEXT',     'deployment.environment resource attribute'),
    ('attrs',        'JSON TEXT','Full resource attributes blob'),
    ('last_seen_ns', 'INTEGER',  'Nanosecond timestamp of most-recent telemetry from this entity'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '2.3  Entity Type Detection')
p('When building EntityRows from a span batch:')
bullet('If resource attrs contain service.name → type = "service"')
bullet('If resource attrs contain k8s.node.name (but no service.name) → type = "host"')
bullet('Duplicate entity-per-batch is resolved by taking the row with the largest last_seen_ns')
bullet('Host entity is skipped when service.name is present to avoid creating both pod-name '
       'and service-name entities in the same flush')

h(2, '2.4  Environment Preservation (Bug Fixed)')
p('The SQL upsert uses CASE logic to prevent newer telemetry from overwriting a '
  'previously-stored deployment.environment when the new record lacks that attribute:')
code('ON CONFLICT(entity_type, entity_id) DO UPDATE SET\n'
     '    environment = CASE\n'
     '        WHEN excluded.environment != \'\' THEN excluded.environment\n'
     '        ELSE entities.environment\n'
     '    END,\n'
     '    attrs        = excluded.attrs,\n'
     '    last_seen_ns = MAX(entities.last_seen_ns, excluded.last_seen_ns)')
p('Without this fix, spans arriving without deployment.environment (e.g. infra metrics) '
  'would silently clear the environment field, causing services to vanish from '
  'environment-filtered UI views.')

h(2, '2.5  Entity-Signal Correlation in the UI')
p('The dashboard polling loop (every 5 s) fetches three resources in parallel and joins them client-side:')
bullet('/v1/query/entities?type=service&env=<env> — list of all service entities')
bullet('/v1/query/topology — service→service edges with call/error counts')
bullet('/v1/query/anomalies?limit=500 — all recent anomalies')
p('The join key is entity_id / source_service / target_service. The UI builds '
  '_anomalyByEntity[entityId] → {signal, severity, cpu} which drives node health '
  'colors, pulse animations, incident panel counts, and RCA entry points.')


# ══════════════════════════════════════════════════════════════════════════════
# 3. TRACE / ERROR FINGERPRINTING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '3. Trace & Error Fingerprinting')

h(2, '3.1  Trace Structural Fingerprinting (buildTraceFP)')
p('The trace fingerprint captures the unique structural shape of a distributed trace: '
  'which services called which other services, and via which operations. '
  'Two traces with identical service-to-service call paths produce the same hash '
  'regardless of individual span durations, status codes, or payload content.')

h(3, 'Algorithm')
bullet('Step 1 — Group spans by trace_id. Skip traces with fewer than 2 spans (single-service, no cross-service edges).')
bullet('Step 2 — Build a span index: SpanID → {SpanID, ParentSpanID, ServiceName, OpName}.')
bullet('Step 3 — Walk every span. For each span whose parent belongs to a DIFFERENT service, '
       'emit a cross-service edge string:\n'
       '        "parentService:parentOp → childService:childOp"')
bullet('Step 4 — Sort the edge list lexicographically (order-independent).')
bullet('Step 5 — MD5-hash the joined edge string. Use the first 16 hex chars as the fingerprint hash.')
bullet('Step 6 — Identify the root span (ParentSpanID is empty or "0000000000000000") → root service + root operation.')

code('edges = ["checkout:POST /checkout→payment:Charge",\n'
     '         "checkout:POST /checkout→inventory:Reserve"]\n'
     'sorted(edges)  # deterministic\n'
     'hash = md5("|".join(edges))[:16]  # "a3f9c1d2e4b87650"')

h(3, 'Noise Filtering')
p('Health-check and infrastructure heartbeat paths are excluded via noisePatterns:')
code('/health, /healthz, /readyz, /livez, /ready, /live,\n'
     '/actuator, /ping, /status, /_health, /api/health,\n'
     '/metrics, /favicon, /eureka/, /v1/agent/, /v1/health/, /v1/catalog/')
p('This prevents health-probe traffic from polluting the fingerprint baseline with '
  'thousands of single-operation "traces".')

h(3, 'Lifecycle: Candidate → Baseline')
p('The fingerprinting worker runs every 5 minutes against the last 5 minutes of spans (up to 5 000):')

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['State', 'Condition'])
for r in [
    ('Bootstrap (silent promotion)',
     'If no baseline exists yet (first run), ALL fingerprints are promoted '
     'directly to baseline without emitting anomalies. Prevents false-positive flood on startup.'),
    ('New candidate',
     'Hash not in baseline, not yet seen. Store as is_baseline=false, occurrence_count=1.'),
    ('Growing candidate',
     'Hash seen in a prior window but not promoted. Increment occurrence_count, update last_seen_at.'),
    ('Promoted to baseline',
     'occurrence_count ≥ 3 AND age ≥ 5 minutes → set is_baseline=true. No anomaly emitted.'),
    ('Baseline — re-observed',
     'Update last_seen_at + increment occurrence_count. No anomaly.'),
    ('trace_drift anomaly',
     'Hash is first seen (isCandidate=false) AND baseline already exists. '
     'Fires ONCE per hash, never on repeated sightings of the same novel path.'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(3, 'Fingerprint Record Schema (TraceFingerprintRow)')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Field', 'Description'])
for r in [
    ('hash',             '16-char hex (MD5 of sorted edge list)'),
    ('root_service',     'Service that owns the root span'),
    ('edge_list',        'JSON array of "src:op→dst:op" strings'),
    ('occurrence_count', 'Total times this fingerprint has been observed'),
    ('first_seen_at',    'Unix seconds when first observed'),
    ('last_seen_at',     'Unix seconds when last observed'),
    ('is_baseline',      'TRUE once promoted; suppresses future anomalies for this hash'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '3.2  Error Signature Fingerprinting (buildErrorSig)')
p('Error signatures capture the unique pattern of a failure: service + error type + '
  'HTTP status + operation name. Two errors from the same service/operation/type '
  'map to the same signature regardless of their message text or stack trace.')

h(3, 'Algorithm')
bullet('Query all spans with status_code = 2 (ERROR) in the last 5-minute window (up to 2 000).')
bullet('For each error span, extract:')
bullet('  error_type  — first of: exception.type, error.kind, error.type span attributes; fallback "error"', indent=1)
bullet('  http_status — first of: http.response.status_code, http.status_code span attributes', indent=1)
bullet('  operation   — span name', indent=1)
bullet('Hash = MD5(service | error_type | http_status | operation)[:16]')

h(3, 'Lifecycle: Candidate → Baseline')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['State', 'Condition'])
for r in [
    ('Bootstrap', 'If no baseline: promote all current signatures silently.'),
    ('New candidate', 'Hash not seen. Store as is_baseline=false.'),
    ('Promoted to baseline', 'occurrence_count ≥ 2 AND age ≥ 30 minutes.'),
    ('New error_signature anomaly',
     'First sighting of a NEW hash (not previously a candidate). Severity: '
     '"warning" for 1–4 occurrences in window, "critical" for ≥ 5.'),
    ('Error spike anomaly',
     'Existing BASELINE signature. current_window_count > BaselineRate × 3. '
     'Severity: always "critical". Score = current/baseline ratio.'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(3, 'Baseline Rate Calculation')
p('When a candidate is promoted to baseline, BaselineRate is computed as:')
code('windows = age_seconds / window_duration_seconds  (min 1)\n'
     'BaselineRate = total_occurrences / windows\n'
     '# i.e. average occurrences per 5-minute window')
p('This rate is used as the denominator for spike detection: '
  'if the current window sees >3× the average, fire a spike anomaly.')


# ══════════════════════════════════════════════════════════════════════════════
# 4. STRUCTURAL DRIFT DETECTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '4. Structural Drift Detection')
p('Structural drift refers to changes in the topology or shape of the system that '
  'are unexpected based on observed history. otel-beacon implements two structural '
  'drift detectors: trace-level and topology-level.')

h(2, '4.1  Trace Drift (trace_drift)')
p('Already described in §3.1. Summary: A "trace_drift" anomaly fires the first time a '
  'previously unseen cross-service call graph fingerprint is observed after a baseline '
  'has been established. It fires once per unique fingerprint hash, suppressed on all '
  'subsequent windows.')

h(3, 'Signal attributes')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Field', 'Value'])
for r in [
    ('signal_type',   '"trace_drift"'),
    ('detector_name', '"Trace Structural Drift"'),
    ('metric_name',   '"trace.fingerprint"'),
    ('value',         '1 (presence indicator)'),
    ('score',         '1'),
    ('severity',      '"warning"'),
    ('description',   '"New call path in <rootSvc>: <edge1>, <edge2> [+N more]"'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '4.2  Call Graph Drift (callgraph_drift)')
p('While trace fingerprinting operates at the trace level (per individual request), '
  'call graph drift operates at the topology level (aggregate edges across all traces). '
  'It detects when a NEW service→service connection appears in the service topology '
  'for the very first time, regardless of which specific trace introduced it.')

h(3, 'Implementation: first_seen_at column')
p('The service_topology table has a first_seen_at INTEGER column (Unix seconds). '
  'The RefreshTopology SQL uses INSERT ... ON CONFLICT DO UPDATE which preserves '
  'first_seen_at on existing edges while setting it on new edges:')
code('INSERT INTO service_topology (..., first_seen_at)\n'
     'SELECT ..., unixepoch() AS first_seen_at\n'
     'FROM spans JOIN spans child ...\n'
     'ON CONFLICT(source_service, target_service) DO UPDATE SET\n'
     '    call_count      = excluded.call_count,\n'
     '    error_count     = excluded.error_count,\n'
     '    avg_duration_ms = excluded.avg_duration_ms,\n'
     '    updated_at      = excluded.updated_at\n'
     '-- first_seen_at NOT updated → timestamp of original appearance preserved')
p('Important: prior versions used INSERT OR REPLACE which deleted + reinserted the row, '
  'destroying the original first_seen_at. The ON CONFLICT DO UPDATE form is correct.')

h(3, 'Detection logic (detectCallGraphDrift)')
bullet('Runs every 2 minutes (aligned with topologyWorker refresh interval).')
bullet('Queries edges where first_seen_at >= now - 180s (3-minute window, slightly wider '
       'than the 2-minute tick to tolerate timing jitter).')
bullet('In-memory driftEmitted map[string]bool keyed by "src→tgt" prevents firing twice '
       'for the same edge within a process lifetime.')
bullet('On first detection of a new edge: emit callgraph_drift anomaly. Never re-fires for '
       'the same edge pair even if the topology refreshes multiple times.')

h(3, 'Signal attributes')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Field', 'Value'])
for r in [
    ('signal_type',   '"callgraph_drift"'),
    ('detector_name', '"call_graph_drift"'),
    ('metric_name',   '"new_topology_edge"'),
    ('value',         'call_count on the new edge'),
    ('score',         '1.0'),
    ('algorithm',     '"structural"'),
    ('severity',      '"warning"'),
    ('description',   '"New call path detected: <src> → <tgt> (first seen, N calls)"'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '4.3  Drift vs. Error: Key Distinction')
p('Drift signals (trace_drift, callgraph_drift) are topology-structural anomalies. '
  'They fire on shape changes, NOT on error conditions. A new edge between two healthy '
  'services will still fire callgraph_drift. This is intentional: unexpected call paths '
  'may indicate misconfiguration, shadow traffic, or undocumented dependencies — '
  'all worth investigating regardless of error rate.')


# ══════════════════════════════════════════════════════════════════════════════
# 5. METRIC ANOMALY DETECTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '5. Metric Anomaly Detection')
p('Metric anomaly detection runs inline in the metricsWorker goroutine. Each metric '
  'data point is passed through a Detector.Check(entity, metricName, value) call '
  'before being enqueued for storage. A non-nil return means an anomaly was detected '
  'and is flushed alongside the metric batch.')

h(2, '5.1  Three Detector Implementations')
p('Configured via the ANOMALY_ALGO environment variable (default: mad):')

h(3, 'Z-Score (zscore)')
p('Classic rolling-window approach. Requires ≥ 30 samples before firing.')
code('z = |value - mean| / stddev\n'
     'fire if z > threshold  (default 3.5)\n'
     'severity = "critical" if z > threshold * 2')
bullet('Window size: 100 samples (configurable via WithAnomalyWindow)')
bullet('15-minute cooldown per entity+metric pair after firing')
bullet('Does not fire if stddev = 0 (constant series)')

h(3, 'MAD — Median Absolute Deviation (mad) [DEFAULT]')
p('Robust to outliers and skewed distributions. Recommended threshold: 3.5.')
code('median  = median(window)\n'
     'MAD     = median(|xi - median| for xi in window)\n'
     'score   = 0.6745 * |value - median| / MAD   # Iglewicz-Hoaglin modified Z-score\n'
     'fire if score > 3.5')
bullet('Window size: 100 samples. Requires ≥ 30 samples.')
bullet('0.6745 normalisation makes the score comparable to Z-scores for normally distributed data.')
bullet('MAD = 0 (all values identical) → no fire. 15-minute cooldown.')

h(3, 'EWMA — Exponentially Weighted Moving Average (ewma)')
p('Adapts quickly to trends; no fixed window needed.')
code('diff     = value - ewma_mean\n'
     'ewma_mean    += alpha * diff\n'
     'ewma_variance = (1-alpha) * (ewma_variance + alpha * diff²)\n'
     'score    = |diff| / sqrt(ewma_variance)\n'
     'fire if score > threshold  (recommended 3.0)')
bullet('Alpha = 0.3 (higher = faster adaptation to trends). Configurable via WithEWMAAlpha.')
bullet('First data point initialises the mean; detection begins from second point.')
bullet('15-minute cooldown. stddev = 0 → no fire.')

h(2, '5.2  Metrics Excluded from Detection')
p('The following metric classes are always skipped (skipDetection = true):')
bullet('Monotonic counters (MetricTypeSum.IsMonotonic=true) — always increasing, '
       'delta analysis is meaningless')
bullet('JVM bookkeeping metrics: jvm.class.count, jvm.class.loaded, jvm.class.unloaded, '
       'jvm.gc.duration, jvm.thread.count, jvm.thread.daemon.count, jvm.cpu.count, '
       'jvm.memory.committed, jvm.memory.limit, process.runtime.jvm.classes.loaded')

h(2, '5.3  Anomaly Row Output')
p('All three detectors produce an AnomalyRow with signal_type = "metric":')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Field', 'Value'])
for r in [
    ('entity_id',     'Entity that emitted the metric'),
    ('signal_type',   '"metric"'),
    ('detector_name', '"Metric Anomaly"'),
    ('metric_name',   'OTLP metric name (e.g. "container.cpu.usage")'),
    ('value',         'The anomalous data point value'),
    ('score',         'Modified Z-score / Z-score / EWMA residual'),
    ('mean',          'Baseline mean (or median for MAD)'),
    ('stddev',        'Baseline std dev (or MAD for MAD)'),
    ('algorithm',     '"zscore" | "mad" | "ewma"'),
    ('severity',      '"critical" if score ≥ 2×threshold, else "warning"'),
    ('detected_at',   'time.Now().UnixNano()'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 6. SPAN RATE DETECTION
# ══════════════════════════════════════════════════════════════════════════════
h(1, '6. Span Rate Detection')
p('The spanRateWorker runs two parallel detectors for span-level error rate and P95 latency:')

h(2, '6.1  Fast Span Detector (30-second interval)')
p('Designed for near-real-time anomaly firing (useful for demo scenarios and rapid on-call response). '
  'Compares the CURRENT 30-second window against the PRIOR 30-second window per service, '
  'using ratio-based thresholds rather than statistical baselines.')

h(3, 'Error Rate Logic')
code('curErr  = errors_in_current_window / spans_in_current_window\n'
     'priorErr = errors_in_prior_window  / spans_in_prior_window  (floor 0.01)\n'
     '         OR 0.02 if prior has < 5 spans\n'
     'ratio   = curErr / priorErr\n'
     'FIRE span_error_rate if:\n'
     '    spans_in_current >= 10\n'
     '    AND errors_in_current >= 3\n'
     '    AND ratio >= 3.0')
p('Severity: always "critical". Score = ratio. Algorithm = "ratio".')

h(3, 'Latency Logic')
code('curP95  = P95(durations_in_current_window)\n'
     'priorP95 = P95(durations_in_prior_window)  (floor 10 ms)\n'
     '          OR 50 ms if prior has < 5 spans\n'
     'FIRE span_latency if:\n'
     '    curP95 >= 500 ms\n'
     '    AND ratio = curP95 / priorP95 >= 2.5')
p('Severity: "critical". Score = ratio. Requires ≥ 10 spans in current window.')

h(2, '6.2  Slow MAD Detector (5-minute interval)')
p('Runs after a 3-minute startup delay to allow a baseline to accumulate. '
  'Uses the configured MAD/Z-score/EWMA detector (§5) against each service\'s '
  'aggregated error rate and P95 over the 5-minute window.')
bullet('Minimum 5 spans per service per window to fire.')
bullet('Feeds error rate and P95 as separate metric streams into the detector '
       'keyed by "span.error_rate" and "span.p95_latency_ms".')
bullet('Inherits the 15-minute cooldown from the underlying detector.')
bullet('Complements the fast detector for sustained (non-spiky) degradation that '
       'the ratio-based fast detector might miss.')


# ══════════════════════════════════════════════════════════════════════════════
# 7. SERVICE TOPOLOGY ENGINE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '7. Service Topology Engine')
p('The service topology is a materialised view of which services call which other services, '
  'aggregated from the last 6 hours of span data. It is stored in the service_topology table '
  'and refreshed every 2 minutes.')

h(2, '7.1  RefreshTopology — SQL Strategy')
p('Two INSERT ... ON CONFLICT DO UPDATE statements cover two edge detection cases:')

h(3, 'Case 1: Parent-child via span parent/child relationship')
code('INSERT INTO service_topology (source_service, target_service,\n'
     '                              call_count, error_count, avg_duration_ms,\n'
     '                              updated_at, first_seen_at)\n'
     'SELECT\n'
     '    parent.entity_id AS source_service,\n'
     '    child.entity_id  AS target_service,\n'
     '    COUNT(*)         AS call_count,\n'
     '    SUM(CASE WHEN child.status_code = 2 THEN 1 ELSE 0 END) AS error_count,\n'
     '    AVG(child.duration_ms) AS avg_duration_ms,\n'
     '    unixepoch()      AS updated_at,\n'
     '    unixepoch()      AS first_seen_at  -- only used for NEW rows\n'
     'FROM spans child\n'
     'JOIN spans parent ON child.parent_span_id = parent.span_id\n'
     'WHERE child.entity_id != parent.entity_id\n'
     '  AND child.start_ns >= unixepoch(\'now\', \'-6 hours\') * 1000000000\n'
     'GROUP BY source_service, target_service\n'
     'ON CONFLICT(source_service, target_service) DO UPDATE SET\n'
     '    call_count      = excluded.call_count,\n'
     '    error_count     = excluded.error_count,\n'
     '    avg_duration_ms = excluded.avg_duration_ms,\n'
     '    updated_at      = excluded.updated_at\n'
     '-- first_seen_at is NOT updated: original timestamp preserved')

p('Error count uses status_code = 2 (OTel STATUS_CODE_ERROR), not HTTP status codes.')

h(3, 'Case 2: Inferred via db.system / messaging.system peer attributes')
p('Spans that call databases, caches, or queues don\'t have a child span from the '
  'other side. otel-beacon infers edges from span attributes:')
code('-- Infers "service → postgres" from db.system=postgresql\n'
     'SELECT\n'
     '    s.entity_id AS source_service,\n'
     '    COALESCE(s.span_attrs->>"$.peer.service",\n'
     '             s.span_attrs->>"$.db.name",\n'
     '             s.span_attrs->>"$.messaging.destination",\n'
     '             s.span_attrs->>"$.db.system") AS target_service,\n'
     '    ...\n'
     'WHERE target_service IS NOT NULL\n'
     '  AND s.entity_id != target_service')

h(2, '7.2  Edge Error Rate & UI Coloring')
p('The UI reads error_count / call_count for each edge and applies CSS classes:')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Error Rate', 'CSS Class', 'Color'])
for r in [
    ('> 10%',   'topo-edge-hot', 'Red #ef4444 (animated pulse)'),
    ('5%–10%',  'topo-edge-err', 'Red #ef4444 (solid)'),
    ('> 2%',    'topo-edge (orange arrowhead)', 'Orange #f97316 arrowhead'),
    ('≤ 2%',    'topo-edge (green arrowhead)',  'Green #4ecfb3 arrowhead'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()
p('The edge fingerprint (used to decide whether to redraw) includes a quantized error bucket '
  '("hot" | "err" | "ok") so threshold crossings trigger a redraw. This prevents the '
  'case where error_count changes on an existing edge but the fingerprint stays identical '
  'and no redraw occurs.')


# ══════════════════════════════════════════════════════════════════════════════
# 8. MISSING SERVICE DETECTION
# ══════════════════════════════════════════════════════════════════════════════
h(1, '8. Missing Service Detection')
p('StartMissingSvcChecker runs every 10 seconds and compares each service entity\'s '
  'last_seen_ns against the current time. If a service has not reported any telemetry '
  'for 45 seconds, a missing_service anomaly is inserted. When the service resumes, '
  'the anomaly is immediately deleted.')

h(2, '8.1  Detection Logic')
code('for each service entity:\n'
     '    silentSec = (now - last_seen_ns) / 1e9\n'
     '    if silentSec > 45:\n'
     '        DELETE existing missing_service anomaly for this entity\n'
     '        INSERT missing_service anomaly (score = silentSec / 45)\n'
     '    else:\n'
     '        DELETE missing_service anomaly (service is back)')
p('The delete-then-insert pattern ensures only one missing_service row per entity at '
  'a time. Without the delete, a 10-second interval would accumulate a new row '
  'every 10 seconds for the entire outage period.')

h(2, '8.2  Denylist')
p('Certain infrastructure component names are excluded because they are expected to '
  'appear as span peer attributes but never emit their own telemetry:')
code('h2, caffeine-cache, config-server, discovery-server,\n'
     'eureka, zipkin, jaeger, prometheus, grafana, loki, otel-collector')

h(2, '8.3  Signal attributes')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Field', 'Value'])
for r in [
    ('signal_type',   '"missing_service"'),
    ('detector_name', '"staleness"'),
    ('value',         'Seconds since last telemetry'),
    ('score',         'silentSec / 45  (>1.0 = definitely missing)'),
    ('algorithm',     '"threshold"'),
    ('severity',      '"critical"'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 9. CROSS-SIGNAL CORRELATION ENGINE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '9. Cross-Signal Correlation Engine (correlator.go)')
p('StartCorrelator runs every 60 seconds. It groups all recent anomalies by entity_id '
  'over a 5-minute rolling window and fires a correlated_incident anomaly when an entity '
  'accumulates 2 or more distinct signal types simultaneously.')

h(2, '9.1  Logic')
code('byEntity = QueryRecentAnomaliesByEntity(ctx, window=300s)\n'
     '\n'
     'for entity, anomalies in byEntity:\n'
     '    sigSet = {a.signal_type for a in anomalies\n'
     '              if a.signal_type != "correlated_incident"}\n'
     '\n'
     '    if len(sigSet) < 2:\n'
     '        continue  # not enough distinct signals\n'
     '\n'
     '    if lastFiredAt[entity] within 5 minutes:\n'
     '        continue  # suppressed\n'
     '\n'
     '    severity = "critical" if any(a.severity == "critical")\n'
     '               else "warning"\n'
     '\n'
     '    emit correlated_incident(entity, signals=sorted(sigSet), severity)')

h(2, '9.2  Parameters')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Constant', 'Value', 'Description'])
for r in [
    ('correlatorWindow',    '5 minutes', 'Look-back window for co-occurrence'),
    ('correlatorInterval',  '60 seconds','How often the correlator runs'),
    ('correlatorMinSignals','2',          'Minimum distinct signal types to fire'),
    ('correlatorSuppress',  '5 minutes', 'Suppression window per entity after firing'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '9.3  Severity Escalation')
p('If any contributing anomaly has severity = "critical", the correlated_incident '
  'inherits "critical" severity. Otherwise it is "warning". This means a '
  'span_error_rate (critical) + trace_drift (warning) = critical correlated_incident.')

h(2, '9.4  Signal attributes')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Field', 'Value'])
for r in [
    ('signal_type',   '"correlated_incident"'),
    ('detector_name', '"cross_signal_correlator"'),
    ('metric_name',   '"correlated_signals"'),
    ('value',         'Number of distinct signal types co-occurring'),
    ('score',         'Same as value'),
    ('algorithm',     '"correlation"'),
    ('description',   '"Correlated incident on <entity>: <sig1> + <sig2> (N signal types co-occurring)"'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 10. ROOT CAUSE ANALYSIS ENGINE
# ══════════════════════════════════════════════════════════════════════════════
h(1, '10. Root Cause Analysis Engine (rca.go)')
p('The RCA endpoint (/v1/rca?entity=<id>&window=<secs>&ts=<unix>) performs a '
  'multi-factor causal analysis for a service at a point in time. The analysis '
  'window defaults to 5 minutes; the incident timestamp defaults to now.')

h(2, '10.1  Data Gathered')
bullet('Focal entity — span health (error rate, P95, avg, total), log error/warn counts, '
       'CPU and memory metrics (container.cpu.usage, container.memory.rss) for the incident window.')
bullet('Baseline — same metrics for the previous equally-sized window (before the incident).')
bullet('Upstream services — all callers of the focal entity (topology edges where target = focal).')
bullet('Downstream services — all callees of the focal entity (topology edges where source = focal).')
bullet('Co-located services — services on the same k8s.node.name as the focal entity.')
bullet('Error signatures — non-baseline error patterns for the focal entity.')
bullet('Trace fingerprints — non-baseline call path fingerprints for the focal entity.')
bullet('Recent anomalies — check for active missing_service anomaly.')

h(2, '10.2  Candidate Cause Ranking (rankCauses)')
p('Each candidate cause receives a confidence score 0–1. Causes are sorted by confidence descending:')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['cause_type', 'Trigger condition', 'Base confidence'])
for r in [
    ('service_down',      'active missing_service anomaly',                '0.99'),
    ('downstream_error',  'downstream service error rate > 5%',            '0.40 + error_rate (cap 0.95)'),
    ('infra_pressure',    'co-located service CPU > 50% or log errors ≥ 10','0.30–0.80'),
    ('upstream_spike',    'upstream span count > 2× focal span count',     '0.35'),
    ('error_signature',   'non-baseline error signature on focal',         '0.60–0.90'),
    ('trace_drift',       'non-baseline trace fingerprint on focal',       '0.55–0.85'),
    ('self_error',        'focal error rate > 5%',                         '0.40 + error_rate (cap 0.80)'),
    ('self_latency',      'focal P95 > 2× baseline P95',                   '0.40–0.85'),
    ('self_infra',        'focal CPU > 70%',                               '0.30–0.75'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '10.3  Temporal Lag Detection')
p('For each neighbor, the RCA engine checks whether degradation preceded the focal incident:')
code('mid = fromNs + (windowNs / 2)\n'
     'early = entityHealth(neighbor, fromNs, mid)\n'
     'if early.error_rate > 10%  OR  early.P95 > full_window_P95 * 1.5:\n'
     '    lag = -(windowSecs / 2) seconds  # "degraded N seconds before focal"\n'
     '    confidence += 0.25  # upstream degradation precedes focal → higher confidence')
p('A negative lag_seconds in the RCA response means the neighbor was degraded '
  'in the first half of the incident window — a strong causal indicator.')

h(2, '10.4  AI Narrative (optional)')
p('If ai=true is passed, the RCA engine calls AWS Bedrock (Claude) with the '
  'RCAResult JSON to generate a natural-language narrative. Falls back gracefully '
  'if Bedrock credentials are unavailable.')


# ══════════════════════════════════════════════════════════════════════════════
# 11. ANOMALY SIGNAL PRIORITY & UI RESOLUTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '11. Signal Priority & UI Rendering')
p('When multiple anomalies exist for the same entity, the UI resolves the "worst" '
  'signal using a priority map (higher = more severe):')

code('sigPrio = {\n'
     '    error_signature:    5,   // new error type → most severe\n'
     '    span_error_rate:    4,   // error rate spike\n'
     '    trace_drift:        3,\n'
     '    correlated_incident:3,\n'
     '    span_latency:       2,\n'
     '    genai_cost_spike:   2,\n'
     '    genai_latency_drift:2,\n'
     '    genai_context_bloat:2,\n'
     '    metric:             1,\n'
     '    callgraph_drift:    1,\n'
     '}')

p('Only anomalies from the last 30 seconds are considered for node health coloring '
  '(using the topoAnomalyCutoffNs = now - 30s filter). This ensures nodes return '
  'to healthy state quickly when anomalies resolve.')

h(2, '11.1  Node Color Mapping')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Signal Type', 'CSS Class', 'Visual'])
for r in [
    ('span_error_rate',    'topo-errorrate',  'Red ring + pulse (fast)'),
    ('error_signature',    'topo-error',      'Orange ring + pulse'),
    ('genai_cost_spike',   'topo-errorrate',  'Red ring + pulse'),
    ('genai_context_bloat','topo-errorrate',  'Red ring + pulse'),
    ('genai_latency_drift','topo-latency',    'Yellow ring + pulse'),
    ('span_latency',       'topo-latency',    'Yellow ring + pulse'),
    ('trace_drift',        'topo-drift',      'Cyan ring + pulse'),
    ('callgraph_drift',    'topo-drift',      'Cyan ring + pulse'),
    ('correlated_incident critical','topo-errorrate','Red ring + pulse'),
    ('correlated_incident warning', 'topo-error',   'Orange ring + pulse'),
    ('missing_service',    'topo-missing',    'Gray ring (dim)'),
    ('metric critical',    'topo-errorrate',  'Red ring + pulse'),
    ('metric warning',     'topo-latency',    'Yellow ring + pulse'),
    ('(none)',             'topo-healthy',    'Green ring (no pulse)'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '11.2  Animation Stability (CSS Class Guarding)')
p('Re-setting the class attribute to the same value restarts CSS animations in all browsers. '
  'To prevent this, updateNodeColors() guards the setAttribute call:')
code('function updateNodeColors() {\n'
     '    svgSel.selectAll(".topo-node").each(function(d) {\n'
     '        const cls = "topo-node " + nodeHealthState(d.id);\n'
     '        if (this.getAttribute("class") !== cls)  // <-- guard\n'
     '            this.setAttribute("class", cls);\n'
     '    });\n'
     '}')
p('Three-tier fingerprinting prevents unnecessary DOM work:')
bullet('Case 1 (fingerprint unchanged) → updateNodeColors() only. No edge churn.')
bullet('Case 2 (node set same, edges changed) → redrawEdgesInPlace() + updateNodeColors(). No SVG teardown.')
bullet('Case 3 (node set changed) → full SVG rebuild. Rare after initial warmup.')


# ══════════════════════════════════════════════════════════════════════════════
# 12. DATA RETENTION & STORAGE
# ══════════════════════════════════════════════════════════════════════════════
h(1, '12. Data Retention & Storage')

h(2, '12.1  Retention Worker')
p('Runs every hour. Deletes rows older than RETENTION_DAYS (default 30) from all tables. '
  'ClickHouse backend is a no-op (uses table-level TTL instead).')

h(2, '12.2  SQLite Schema (key tables)')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Table', 'Purpose'])
for r in [
    ('spans',              'Raw span rows. Queried by fingerprint, rate, and RCA workers.'),
    ('metrics',            'Raw metric data points.'),
    ('logs',               'Log records.'),
    ('anomalies',          'All fired anomalies. entity_id + signal_type + detected_at indexed.'),
    ('entities',           'Discovered service/host entities. PK = (entity_type, entity_id).'),
    ('service_topology',   'Service dependency edges. PK = (source_service, target_service). '
                           'Columns: call_count, error_count, avg_duration_ms, updated_at, first_seen_at.'),
    ('trace_fingerprints', 'Trace structural fingerprints. PK = hash.'),
    ('error_signatures',   'Error pattern signatures. PK = hash.'),
    ('genai_spans',        'GenAI-specific span data (model, tokens, cost, eval results).'),
    ('eval_results',       'LLM eval results per span.'),
    ('guardrail_events',   'PII/injection/toxicity check results.'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '12.3  Reset')
p('POST /v1/scenarios/reset calls ResetSimulationData which truncates:')
code('spans, metrics, logs, entities, anomalies,\n'
     'error_signatures, trace_fingerprints, service_topology')
p('This is a full clean-slate reset for demo/scenario purposes. All in-memory '
  'detector windows and driftEmitted maps persist in the Go process (requires restart '
  'for a complete memory reset).')


# ══════════════════════════════════════════════════════════════════════════════
# 13. GENAI OBSERVABILITY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(1, '13. GenAI Observability Layer')
p('Spans carrying gen_ai.* attributes (OTel GenAI semantic conventions) are '
  'intercepted in InsertTraces and routed to a parallel genaiWorker pipeline '
  'in addition to the normal spans table.')

h(2, '13.1  GenAI Span Detection')
code('func isGenAISpan(sp ptrace.Span) bool {\n'
     '    // True if span name starts with "gen_ai." or has gen_ai.* attributes\n'
     '}')

h(2, '13.2  GenAI Anomaly Detectors')
p('The genaiWorker computes per-model cost, latency, and token metrics and '
  'feeds them into the same MAD/Z-score/EWMA detector for anomaly detection:')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_header(tbl, ['Signal Type', 'Trigger'])
for r in [
    ('genai_cost_spike',    'Cost-per-request spike above baseline (MAD score > threshold)'),
    ('genai_latency_drift', 'LLM response latency spike'),
    ('genai_context_bloat', 'Token count (prompt + completion) spike'),
]:
    add_table_row(tbl, r)
doc.add_paragraph()

h(2, '13.3  LLM-as-Judge Eval Worker (StartEvalWorker)')
p('Async background worker that drains genaiEvalCh and calls AWS Bedrock to '
  'evaluate each GenAI span across four dimensions:')
bullet('Hallucination proxy (answer-vs-context consistency check)')
bullet('Coherence (logical structure of the response)')
bullet('Relevance (response relevance to the prompt)')
bullet('Toxicity (harmful content detection)')
p('Falls back to heuristic scoring if Bedrock returns AccessDenied or is unreachable.')


# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX: SIGNAL TYPE REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h(1, 'Appendix A: Signal Type Quick Reference')

tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
table_header(tbl, ['signal_type', 'Source', 'Interval', 'Severity', 'Cooldown / Suppression'])
rows = [
    ('metric',             'metricsWorker (inline)', '500ms flush', 'warning / critical (2× threshold)', '15 min per entity+metric'),
    ('span_error_rate',    'Fast: spanRateWorker\nSlow: spanRateWorker', 'Fast: 30s\nSlow: 5min', 'critical (fast)\nwarning/critical (slow)', 'Fast: none\nSlow: 15 min'),
    ('span_latency',       'Same as above',           'Same',        'critical (fast)\nwarning/critical (slow)', 'Same'),
    ('trace_drift',        'fingerprintWorker',       '5 min',       'warning',  'Fire once per hash lifetime'),
    ('error_signature',    'errorSignatureWorker',    '5 min',       'warning/critical', 'Fire once per hash (new)\nFire on spike (baseline)'),
    ('callgraph_drift',    'callGraphDriftWorker',    '2 min',       'warning',  'Fire once per src→tgt (in-memory dedup)'),
    ('missing_service',    'StartMissingSvcChecker',  '10 s',        'critical', 'Delete+reinsert (1 row per entity)'),
    ('correlated_incident','StartCorrelator',         '60 s',        'warning/critical', '5 min per entity'),
    ('genai_cost_spike',   'genaiWorker',             'per span',    'warning/critical', '15 min (MAD cooldown)'),
    ('genai_latency_drift','genaiWorker',             'per span',    'warning/critical', '15 min'),
    ('genai_context_bloat','genaiWorker',             'per span',    'warning/critical', '15 min'),
]
for r in rows:
    add_table_row(tbl, r)

doc.add_paragraph()
h(1, 'Appendix B: Configuration Environment Variables')

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
table_header(tbl, ['Variable', 'Default', 'Description'])
for r in [
    ('DB_DRIVER',          'sqlite',       '"sqlite" or "clickhouse"'),
    ('DB_DSN',             'otel.db',      'SQLite file path or ClickHouse DSN'),
    ('ANOMALY_ALGO',       'mad',          '"mad" | "zscore" | "ewma"'),
    ('ANOMALY_THRESHOLD',  '3.5',          'Detector threshold (modified Z-score for MAD)'),
    ('RETENTION_DAYS',     '30',           'Delete data older than N days'),
    ('AUTH_TOKEN',         '(none)',       'Bearer token for OTLP endpoints (optional)'),
    ('HTTP_ADDR',          ':4318',        'OTLP/HTTP listen address'),
    ('GRPC_ADDR',          ':4317',        'OTLP/gRPC listen address'),
    ('ADMIN_ADDR',         ':8080',        'Admin + query API + UI listen address'),
    ('TLS_CERT_FILE',      '(none)',       'TLS certificate path (enables TLS when set)'),
    ('TLS_KEY_FILE',       '(none)',       'TLS key path'),
]:
    add_table_row(tbl, r)

doc.add_paragraph()

# Footer note
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_footer.add_run(
    f'otel-beacon Engineering Reference  ·  Generated {datetime.date.today().isoformat()}  ·  INTERNAL')
r.font.size = Pt(8)
r.font.color.rgb = GRAY

# ─── Save ──────────────────────────────────────────────────────────────────────
out_path = 'docs/otel-beacon-anomaly-detection-engine.docx'
doc.save(out_path)
print(f'Written: {out_path}')
